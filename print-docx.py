#!/usr/bin/env python3

""" print-docx.py at https://github.com/wilsonmar/python-samples/blob/main/print-docx.py
STATUS: NOT Working: docx import not in conda
"v006 + get file size :print-docx.py"

Based on a specified (hard coded) folder/file path,
this program creates a .docx format Microsoft Word file, then
converts it to a .pdf file for generic printing on the default printer.

This is a local solution instead of having the 
printer company send pages over the internet:
https://www.youtube.com/watch?v=53nOTAc0Oy8
That EpsonConnect technique prints out a cover sheet as well the file
if you don't uncheck "print email body" under Email Print settings.

Before running this program:
brew install miniconda
conda create -n py312
conda install -c conda-forge python=3.12
pip install docx docx2pdf
conda activate py312
chmod +x print-docx.py
./print-docx.py
"""
import docx
from docx2pdf import convert
# import schedule

# built-in modules:
import os
import datetime
import subprocess   # to run CLI commands

# Globals:
MY_PRINTER = "EPSON_ET_2850_Series"
OVERWRITE_FILE_CREATED = False
# TODO: Enable specification of path in a CLI attribute:
# TODO: docx_file_path = fr"{user_home_path}/{docx_file_name}"
docx_file_name = "print-docx.docx"
docx_file_path = fr"/Users/johndoe/github-wilsonmar/python-samples/{docx_file_name}"
pdf_file_name = "print-docx.pdf"
pdf_file_path = fr"/Users/johndoe/github-wilsonmar/python-samples/{pdf_file_name}"
LIST_PRINTERS = False
SHOW_FILE_CONTENTS = False
PRINTER_NAME = "EPSON_ET_2850_Series"
REMOVE_PDF_FILE_CREATED = False


def get_file_datetime(file_path):
    # Get the file creation timestamp:
    creation_time = os.path.getctime(docx_file_path)
    # Convert the timestamp to a datetime object:
    creation_date = datetime.datetime.fromtimestamp(creation_time)
    # Print the creation date in a human-readable format:
    return creation_date.strftime('%Y-%m-%d %H:%M:%S')


def get_file_size(file_path):
    """ Get the file size in bytes """
    try:
        file_size = os.path.getsize(file_path)
        return file_size
    except FileNotFoundError:
        print(f"Error: File '{file_path}' not found.")
        return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None


def create_word_document(file_path):
    # Create a new Document:
    doc = Document()

    # Add a title:
    # doc.add_heading('Ink Colors', 0)

    # Define colors and their names:
    colors = [  # flake8: noqa: F821 undefined name 'Document'
        ('Yellow', RGBColor(255, 255, 0)),
        ('Magenta', RGBColor(255, 0, 255)),
        ('Cyan', RGBColor(0, 255, 255)),
        ('Black', RGBColor(0, 0, 0))
    ]

    # TODO: Remove "Ink Colors" heading. This can be ignored for now.

    # Add a paragraph for each color using VBA coding:
    for color_name, color_rgb in colors:
        # TODO: Randomize location on page:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("---")
        font = run.font
        font.color.rgb = color_rgb

    # Save the document:
    doc.save(file_path)
    # print(f"Document '{file_path}' created.")


def list_printers():
    if os.name == 'nt':  # For Windows:
        # See https://www.blog.pythonlibrary.org/2010/02/14/python-windows-and-printers/
        # os.startfile(pdf_file_path, "print")
        print("*** FIXME: Add coding to print pdf_file_path on Windows.")
    else:  # For POSIX-based systems (Linux, macOS):
        try:
            result = subprocess.run(['lpstat', '-p'],
                                    capture_output=True,
                                    text=True,
                                    check=True)
            # Split the output into separate lines:
            printer_lines = result.stdout.strip().split('\n')
            # Process each line to extract printer names:
            printers = []
            for line in printer_lines:
                # Example: EPSON_ET_2850_Series is idle.  enabled since Sat Oct 19 09:32:51 2024
                if line.startswith('printer'):
                    # Extract the printer name (second word in the line)
                    printer_name = line.split()[1]
                    printers.append(printer_name)
            return printers
        except subprocess.CalledProcessError:
            print("Error: Unable to retrieve printer list.")
            return []

# See https://www.hexnode.com/mobile-device-management/help/script-to-set-up-printers-on-macos-devices/


def print_docx(file_path, printer=""):

    # TODO: Send to specified printer

    # Read the content of the .docx file
    doc = docx.Document(file_path)

    if SHOW_FILE_CONTENTS:
        print("Content of the .docx file:")
        for paragraph in doc.paragraphs:
            print(paragraph.text)

    # Convert .docx to .pdf
    pdf_path = file_path.rsplit('.', 1)[0] + '.pdf'
    convert(file_path, pdf_path)

    # Print the PDF file
    if os.name == 'nt':  # For Windows
        os.startfile(pdf_path, "print")
    else:  # For Unix-based systems (Linux, macOS)
        subprocess.run(["lpr", pdf_path])
        # 100%|████...████| 1/1 [00:21<00:00, 21.65s/it]

    print(f"*** {pdf_path} sent to the default? printer.")


if __name__ == "__main__":

    if os.path.exists(docx_file_path):
        # Get the file creation timestamp:
        creation_time = os.path.getctime(docx_file_path)
        # Convert the timestamp to a datetime object:
        creation_date = datetime.datetime.fromtimestamp(creation_time)
        if OVERWRITE_FILE_CREATED:
            print(f"*** INFO: {docx_file_path} File not found. Creating...")
            create_word_document(docx_file_path)
        else:
            # Print the creation date in a human-readable format:
            print(f"*** INFO: {docx_file_path} found...")
            print("*** INFO: with created date/timestamp "+get_file_datetime(docx_file_path))
            print("*** INFO: file size bytes: "+get_file_size(docx_file_path))
            # Don't create_word_document().

    if LIST_PRINTERS:
        available_printers = list_printers()
        if available_printers:
            print("*** Available printers using lpstat -p command:")
            for printer in available_printers:
                print(f"- {printer}")
        else:
            print("*** CRTICAL: No printers found. Exiting...")
            exit()

    # TODO: Print to MY_PRINTER retrieved from env file.
    print_docx(docx_file_path)

    # Clean up: remove the temporary PDF file
    if REMOVE_PDF_FILE_CREATED:
        os.remove(pdf_file_path)


    # Schedule the printing task to run daily at a specific time (e.g., 9:00 AM)
    # schedule.every().day.at("09:00").do(print_sheet)
    # Keep the script running:
    #while True:
    #    schedule.run_pending()
    #    time.sleep(1)
