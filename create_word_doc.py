#!/usr/bin/env python3

"""create_word_doc.py at https://github.com/wilsonmar/python-samples/blob/main/create_word_doc.py
This program creates a file in Microsoft Word .docx format.

// SPDX-License-Identifier: MIT
CURRENT STATUS: UNDER TEST.

git commit -m"v001 + new :create_word_doc.py"

Tested on macOS 24.1.0 using Python 3.12.8
flake8  E501 line too long, E222 multiple spaces after operator

"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_formatted_doc():
    # Create a new Document
    doc = Document()

    # Add title
    title = doc.add_heading('Python Samples Repository Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Read the content from our text file
    with open('repo-analysis.docx', 'r') as file:
        content = file.readlines()

    # Process each line
    for line in content:
        if line.strip().startswith('---'):
            continue
        elif line.strip().endswith(':'):
            # Section headers
            p = doc.add_heading(line.strip(), level=1)
        elif line.strip().startswith('-'):
            # Bullet points
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        else:
            # Regular text
            p = doc.add_paragraph(line.strip())

    # Save the document
    doc.save('Repository_Analysis.docx')

if __name__ == '__main__':
    create_formatted_doc()
